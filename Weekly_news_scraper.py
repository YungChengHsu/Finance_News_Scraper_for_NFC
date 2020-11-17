#!/usr/bin/env python
# coding: utf-8

# In[1]:


from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

import requests
import time


# In[2]:


# open a new .docx file
document = Document()

document.add_heading('產業新聞', 0).bold = True

#Yahoo news
yh_page = requests.get('https://tw.stock.yahoo.com/intl-markets')
yh_soup = BeautifulSoup(yh_page.text, 'html.parser')
yh_news_links = yh_soup.find_all('a', 
                         class_='Fw(b) Fz(20px) Lh(23px) LineClamp(2,46px) C($c-primary-text)! C($c-active-text)!:h mega-item-header-link Td(n) C(#0078ff):h C(#000) LineClamp(2,46px) LineClamp(2,38px)--sm1024 not-isInStreamVideoEnabled', 
                         href=True)
print(yh_news_links)
yh_num_of_news = len(yh_news_links)

print('Processing Yahoo...')

for i in range(0, yh_num_of_news):
    news_html = requests.get(yh_news_links[i]['href'])
    print(news_html)
    news_soup = BeautifulSoup(news_html.text, 'html.parser')
    
    news_title = news_soup.find_all('h1')[0].get_text()
    
    h = document.add_heading('', level=1)
    h.bold = True
    h.add_run(news_title).font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    news_body = news_soup.find_all('p', class_='canvas-atom canvas-text Mb(1.0em) Mb(0)--sm Mt(0.8em)--sm')
    news_body_p_len = len(news_body)
    for j in range(0, news_body_p_len):
        c = document.add_paragraph(news_body[j].get_text())
        if len(news_body[j].get_text()) > 0 and news_body[j].get_text()[0] == '（' and news_body[j].get_text()[-1] == '）':
            break
    
    c.add_run('\n')

#36Kr
kr_page = requests.get('https://36kr.com/search/articles/8点1氪')
kr_soup = BeautifulSoup(kr_page.content, 'html.parser')
#kr_news_header = kr_soup.find_all('a', class_="article-item-title weight-bold")
kr_news_links = kr_soup.find_all('a', class_="article-item-description ellipsis-2", href=True)
kr_time = kr_soup.find_all('span', class_="kr-flow-bar-time")
kr_num_of_news = len(kr_time)

for i in range(0, kr_num_of_news):
    time_str = kr_time[i].get_text()
    
    # Calculate the time stamp to make sure we only scrape the news within a week
    if time_str != '昨天' and ('前' not in time_str):
        news_time_stamp = time.mktime(datetime.strptime(time_str, "%Y-%m-%d").timetuple())
        now_time_stamp = datetime.timestamp(datetime.now())
        
        if now_time_stamp - news_time_stamp > 604800.0:
            break
        
    news_html = requests.get('https://36kr.com' + kr_news_links[i]['href'])
    news_soup = BeautifulSoup(news_html.text, 'html.parser')
    news_body = news_soup.select('div.common-width.content.articleDetailContent.kr-rich-text-wrapper p')
    
    if len(news_body) == 0:
        continue
    
    print('Processing 36Kr...')
            
    news_body_list = []
    for body in news_body:
        content = body.get_text()
        if content != '' and ('&amp;' not in content):
            news_body_list.append(content)
    
    print(news_body_list)
    
    news_body_len = len(news_body_list)
    
    # in the list, index numbers of even number will be the news title and index numbers of odd number would be the news body
    for j in range(0, news_body_len):
        if j % 2 == 0:
            h = document.add_heading('', level=1)
            h.bold = True
            h.add_run(news_body_list[j]).font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            c = document.add_paragraph(news_body_list[j])
        c.add_run('\n')
    
document.save('產業新聞20201102.docx')

